"""Single channel for Phase 1. Restricted to TELEGRAM_OWNER_ID so this stays
a personal assistant, not an open bot, until Phase 3's multi-user work.
"""
import asyncio
import os
import re

from telegram import InlineKeyboardButton, InlineKeyboardMarkup, Update
from telegram.constants import ChatAction
from telegram.ext import (
    Application,
    CallbackQueryHandler,
    CommandHandler,
    ContextTypes,
    JobQueue,
    MessageHandler,
    filters,
)

from kyraan.agents import orchestrator
from kyraan.control_plane.dnd import local_now
from kyraan.control_plane.logging_setup import get_logger, log_event
from kyraan.triggers import briefs, scheduler

def _uploader(update) -> str:
    """The registry person who sent a file: the owner, or the household
    member whose chat this is (documents get a who-sent-it link)."""
    try:
        from kyraan.store import persons
        if update.effective_user is not None and update.effective_user.id == _owner_id():
            return "owner"
        return persons.person_id_any_stage(update.effective_chat.id) or ""
    except Exception:
        return ""


async def _enrich_and_tell(update, doc_id) -> None:
    """The saver engine's second pass runs after the receipt (the local
    model takes seconds): what the document is, whom it concerns, its
    kin — sent as a follow-up when there is something to say."""
    import asyncio as _aio
    if not doc_id:
        return
    try:
        from kyraan.store import documents as _d
        links = await _aio.to_thread(_d.enrich, doc_id)
        note = _d.receipt_line(links).strip()
        if note:
            await update.message.reply_text(note, do_quote=True)
    except Exception as exc:
        logger.warning("document enrich failed: %s", exc)


def _private_note(doc_id) -> str:
    """Tax papers, statements, medical files stay local-only: the owner
    should know the cloud never sees them and the local model answers."""
    try:
        from kyraan.store import documents as _d
        if doc_id and _d.exposure_of(doc_id) == "local_only":
            return " 🔒 Kept private: only the local model reads it, never the cloud."
    except Exception:
        pass
    return ""


logger = get_logger("telegram_bot")


def _owner_id() -> int:
    return int(os.environ["TELEGRAM_OWNER_ID"])


def _owner_private(update: Update) -> bool:
    """Owner id AND a private chat. The id check alone let the owner's
    messages in any group trigger replies INTO that group — personal
    data and confirm flows in front of whoever else is there (external
    review P1). Group/channel updates are ignored outright."""
    return (update.effective_user is not None
            and update.effective_user.id == _owner_id()
            and update.effective_chat is not None
            and getattr(update.effective_chat, "type", None) == "private")


def _authorized(update: Update) -> str | None:
    """P3.5a: who this update is from — 'owner' (env check, unchanged,
    never touches a store), an enrolled person id at stage >=
    read_mostly in a PRIVATE chat, or None (rejected exactly as before
    Phase 3). Media/voice/location handlers stay owner-only until
    P3.5b/c scope tools and visibility for other stages."""
    if _owner_private(update):
        return "owner"
    if (update.effective_chat is None
            or getattr(update.effective_chat, "type", None) != "private"
            or update.effective_user is None
            or update.effective_user.id != update.effective_chat.id):
        return None
    from kyraan.store import persons
    row = persons.person_for_chat(update.effective_chat.id)
    return row[0] if row else None


def _plain(text: str) -> str:
    """Models emit markdown bold/italic; replies are sent without
    parse_mode (entity-parse errors on unbalanced markers would eat whole
    messages), so Telegram shows the raw asterisks — strip them."""
    return text.replace("**", "").replace("__", "")


async def _typing_loop(bot, chat_id: int) -> None:
    """Keep the "Kyraan is typing…" indicator alive while a reply is being
    produced — Telegram expires a chat action after ~5s, and model + tool
    time regularly exceeds that. Cancelled the moment the reply is ready;
    the indicator also tells the owner their message was received (bots
    can't mark messages "read" — that's a Business-API-only feature)."""
    try:
        while True:
            await bot.send_chat_action(chat_id=chat_id, action=ChatAction.TYPING)
            await asyncio.sleep(4)
    except asyncio.CancelledError:
        pass


def _confirm_keyboard(chat_id: int) -> InlineKeyboardMarkup | None:
    """Yes/No buttons whenever a confirmation is pending for this chat —
    tap instead of typing, and a tap is unambiguous in a way a later
    typed "yes" never is."""
    if chat_id not in orchestrator._pending_confirmations:
        return None
    nonce = orchestrator._confirmation_nonce.get(chat_id, "")
    return InlineKeyboardMarkup([[
        InlineKeyboardButton("✅ Yes", callback_data=f"kyraan_yes:{nonce}"),
        InlineKeyboardButton("❌ No", callback_data=f"kyraan_no:{nonce}"),
    ]])


async def _on_callback(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    query = update.callback_query
    if _authorized(update) is None:
        await query.answer()
        return
    await query.answer()
    action, _, nonce = (query.data or "").partition(":")
    chat_id = update.effective_chat.id
    word = "yes" if action == "kyraan_yes" else "no"
    # Remove the buttons from the ask so a decided confirmation can't be
    # tapped twice.
    try:
        await query.edit_message_reply_markup(reply_markup=None)
    except Exception:
        pass  # message may be old/edited — the confirm flow still decides
    # Nonce validated INSIDE the per-chat lock (security rounds 2-3, P1):
    # checked outside, a concurrent message could replace the pending
    # action while this old button waited on the lock — and the stale Yes
    # would then confirm the NEWER action.
    async with _lock_for(chat_id):
        if nonce != orchestrator.current_confirmation_nonce(chat_id):
            await context.bot.send_message(
                chat_id=chat_id,
                text="That button belongs to an earlier ask and is no longer "
                     "active — reply to the latest confirmation instead.")
            return
        from kyraan.control_plane import kernel as _kernel
        stage_token = _kernel.set_viewer(*_viewer_for(update))
        try:
            reply = await orchestrator.handle_message(chat_id, word)
        finally:
            _kernel.reset_viewer_stage(stage_token)
    # where it was processed, on every reply (owner 2026-09-03)
    reply = f"{reply}\n\n{orchestrator.processing_marker(chat_id)}"

    async def _send_pieces():
        for piece in _pieces(_plain(reply)):
            await context.bot.send_message(chat_id=chat_id, text=piece)

    await _deliver(chat_id, _send_pieces, reply)


_TG_MAX = 4000   # Telegram's hard limit is 4096; keep headroom for markers


def _pieces(text: str) -> list:
    """Telegram refuses messages over 4096 chars (a long PDF read, a full
    listing) — split on paragraph, then line, then hard (review 2026-09-03)."""
    if len(text) <= _TG_MAX:
        return [text]
    out, cur = [], ""
    for para in text.split("\n\n"):
        while len(para) > _TG_MAX:
            cut = para.rfind("\n", 0, _TG_MAX)
            cut = cut if cut > 0 else _TG_MAX
            if cur:
                out.append(cur); cur = ""
            out.append(para[:cut]); para = para[cut:].lstrip("\n")
        if len(cur) + len(para) + 2 > _TG_MAX:
            out.append(cur); cur = para
        else:
            cur = f"{cur}\n\n{para}" if cur else para
    if cur:
        out.append(cur)
    return [p for p in out if p]


async def _deliver(chat_id: int, send, reply_preview: str) -> bool:
    """Delivery truth (CommitKernel-lite, 2026-08-28): an action may have
    EXECUTED even when Telegram drops the receipt — the turn's execution
    status and its delivery status are different facts. One retry, then
    an event that records the divergence with the undelivered text, so
    "the AC switched but the owner never saw the receipt" is a greppable
    fact instead of a mystery."""
    from kyraan.control_plane.logging_setup import log_event
    from telegram.error import BadRequest, TimedOut
    try:
        await send()
        return True
    except BadRequest as exc:
        # the same message will fail the same way (too long, bad markup)
        log_event("reply_delivery_failed", chat_id=chat_id, error=str(exc)[:120],
                  undelivered=reply_preview[:300])
        return False
    except TimedOut as exc:
        # Telegram may have accepted it — a retry risks a duplicate
        log_event("reply_delivery_uncertain", chat_id=chat_id, error=str(exc)[:120])
        return True
    except Exception as exc:
        log_event("reply_delivery_retry", chat_id=chat_id,
                  error=str(exc)[:120])
        await asyncio.sleep(1.5)
        try:
            await send()
            return True
        except Exception as exc2:
            log_event("reply_delivery_failed", chat_id=chat_id,
                      error=str(exc2)[:120],
                      undelivered=reply_preview[:300])
            return False


# Burst coalescing, modeled on how two humans chat: B watches A's typing
# indicator and replies only when A's thought looks COMPLETE; if a new
# message lands while B is mid-reply, B stops, reads it, and rethinks.
# Telegram never sends bots the typing indicator, so both halves are
# inferred: thought-completeness from the message's shape
# (orchestrator.thought_open), and "stop and rethink" from a fragment
# arriving while a reply is still being planned (supersede — the draft is
# retracted and re-planned with the full thought). Requires
# concurrent_updates so later fragments can join while a window is open;
# a per-chat lock keeps actual processing strictly serialized.
_BURST_WINDOW_S = 2.5           # typing a follow-up message takes 2-5s
_BURST_MAX_WAIT_S = 8.0
_FRAGMENT_EXTRA_WAIT_S = 10.0   # an open thought almost certainly has
                                # more coming — wait patiently for it
_burst_buffers: dict = {}
_burst_flushing: set = set()
_burst_superseded: dict = {}
_chat_locks: dict = {}


def _lock_for(chat_id: int) -> asyncio.Lock:
    if chat_id not in _chat_locks:
        _chat_locks[chat_id] = asyncio.Lock()
    return _chat_locks[chat_id]


_courtesy_sent: dict = {}  # chat_id -> local date of the one daily line


async def _on_message(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    person = _authorized(update)
    if person is None:
        logger.warning("Ignored unauthorized update (user=%s chat=%s)",
                       update.effective_user, update.effective_chat)
        # KNOWN-but-unenrolled (stage 'none' with this chat_id recorded):
        # a family member gets one courtesy line a day instead of pure
        # silence (live 2026-08-27: Ruma messaged from her own phone and
        # concluded the bot was broken). Strangers still get nothing.
        try:
            chat_id = update.effective_chat.id
            from kyraan.store import pg as _pg
            with _pg.connection() as conn:
                row = conn.execute(
                    "SELECT id FROM person WHERE chat_id = %s AND stage = 'none'",
                    (chat_id,)).fetchone()
            today = local_now().date().isoformat()
            if row and _courtesy_sent.get(chat_id) != today:
                _courtesy_sent[chat_id] = today
                await update.message.reply_text(
                    "Hi — I'm Manab's personal assistant and can only talk "
                    "to people he has enrolled. Ask him to switch you on!")
        except Exception:
            pass  # a courtesy must never error
        return
    if person != "owner":
        # An enrolled person gets the TEXT pipeline; face enrollment is
        # a biometric write and stays an owner ceremony (P3.5a scope —
        # tool/visibility scoping per stage lands in P3.5b/c).
        await _ingest(update, context, update.message.text or "")
        return

    from kyraan.agents import faces, orchestrator
    name = faces.enroll_from_text(update.message.text or "")
    if name:
        chat_id = update.effective_chat.id
        stashed = faces.recent_photo(chat_id)
        if stashed is not None and faces.face_count(stashed) == 1:
            reply = await _enroll_face_gated(chat_id, name, stashed)
            orchestrator.record_exchange(chat_id, update.message.text or "", reply)
            await update.message.reply_text(
                _plain(reply), do_quote=True,
                reply_markup=_confirm_keyboard(chat_id))
            return
        # No recent photo: fall through — it may be an ordinary memory
        # statement, and the normal pipeline handles those.

    await _ingest(update, context, update.message.text or "")


def _viewer_stage_for(update: Update) -> str:
    """P3.5b: the stage the kernel scopes this turn to. Owner (env check)
    is unscoped; an enrolled chat gets its enrolled stage; anything else
    is 'none' (belt — such updates were already rejected upstream)."""
    if _owner_private(update):
        return "owner"
    from kyraan.store import persons
    row = persons.person_for_chat(update.effective_chat.id)
    return row[1] if row else "none"


async def _media_admitted(update: Update, capability: str) -> bool:
    """Media joined the capability system (owner, 2026-08-28: "ruma
    trying to upload images but unable" — media was a silent hardcoded
    owner-only wall outside every access model). Owner: always. An
    enrolled viewer: their stage toolset ∪ the owner's individual
    grants ("give ruma media.photo"). Denied enrolled viewers get an
    honest line instead of silence; strangers still get nothing."""
    if _owner_private(update):
        return True
    from kyraan.control_plane import kernel
    person, stage = _viewer_for(update)
    if not person or stage not in ("read_mostly", "full"):
        return False  # not admitted to chat at all — stay silent
    token = kernel.set_viewer(person, stage)
    try:
        if kernel.stage_allows(capability):
            return True
    finally:
        kernel.reset_viewer_stage(token)
    display = person.replace("_", " ").title()
    await update.message.reply_text(
        f"{display}, you don't have {capability.split('.', 1)[1]} access "
        "here — ask Maan to grant it.", do_quote=True)
    return False


def _viewer_for(update: Update) -> tuple:
    """(person_id, stage) — BOTH, always (2026-08-28: setting stage only
    left the viewer person empty, a fail-open default turned empty into
    'owner', and Ruma's first enrolled chat called HER Maan and told her
    she was the owner). Identity is resolved for any registered chat
    regardless of stage; an unregistered chat is ('', 'none')."""
    if _owner_private(update):
        return ("owner", "owner")
    from kyraan.store import persons
    person = persons.person_id_any_stage(update.effective_chat.id) or ""
    return (person, _viewer_stage_for(update))


async def _ingest(update: Update, context: ContextTypes.DEFAULT_TYPE, text: str) -> None:
    """Everything after the ownership gate — shared by typed messages and
    transcribed voice notes (which ARE text by the time they get here)."""
    from kyraan.control_plane import kernel
    stage_token = kernel.set_viewer(*_viewer_for(update))
    try:
        await _ingest_inner(update, context, text)
    finally:
        kernel.reset_viewer_stage(stage_token)


async def _ingest_inner(update: Update, context: ContextTypes.DEFAULT_TYPE, text: str) -> None:
    chat_id = update.effective_chat.id
    _burst_buffers.setdefault(chat_id, []).append((update.message, text))
    if chat_id in _burst_flushing:
        # A window is open or a reply is being composed. The new fragment
        # supersedes any draft: composition checks the event at its safe
        # point (before anything with side effects runs) and starts over
        # with the full thought.
        event = _burst_superseded.get(chat_id)
        if event is not None:
            event.set()
        return
    _burst_flushing.add(chat_id)
    # Typing starts the moment the message lands — it doubles as the
    # "seen" receipt (bots can't mark messages read). Starting it only at
    # composition made it invisible: the fast frontier answers in ~2s and
    # the indicator never got long enough on screen to render.
    typing = asyncio.create_task(_typing_loop(context.bot, chat_id))
    try:
        while True:
            # Gather: wait for quiet, patiently while the last message
            # reads as an unfinished thought — the substitute for watching
            # the typing indicator, which Telegram never sends bots.
            waited = 0.0
            last_text = ""
            while waited < _BURST_MAX_WAIT_S * 3:  # hard cap ~20s+
                buffered = _burst_buffers.get(chat_id, [])
                seen = len(buffered)
                last_text = buffered[-1][1] if buffered else ""
                quiet = _BURST_WINDOW_S
                if seen > 1 or orchestrator.thought_open(last_text):
                    quiet = _BURST_WINDOW_S * 2
                await asyncio.sleep(quiet)
                waited += quiet
                if len(_burst_buffers.get(chat_id, [])) == seen:
                    if orchestrator.thought_open(last_text) and waited < _FRAGMENT_EXTRA_WAIT_S:
                        continue  # thought still open — keep waiting
                    break  # quiet and complete — reply now
            # The event exists BEFORE the buffer is popped, so a fragment
            # arriving in between lands in this pop, and one arriving
            # after it sets the event — no gap either way.
            event = _burst_superseded[chat_id] = asyncio.Event()
            fragments = _burst_buffers.pop(chat_id, [])
            if not fragments:
                return
            try:
                async with _lock_for(chat_id):
                    # The burst is evaluated TOGETHER and answered as ONE
                    # composed reply, quoted onto the message it covers.
                    results = await orchestrator.handle_burst(
                        chat_id, [text for _, text in fragments], superseded=event
                    )
            except orchestrator.BurstSuperseded:
                # The user kept typing while the reply was being planned —
                # nothing ran yet, so retract the draft, fold these
                # fragments back in front of the newcomers, and re-read
                # the whole thought (the typing indicator stays on: Kyraan
                # is still working on a reply).
                _burst_buffers[chat_id] = fragments + _burst_buffers.get(chat_id, [])
                continue
            for position, (idx, reply) in enumerate(results):
                source = fragments[min(idx, len(fragments) - 1)][0]
                markup = _confirm_keyboard(chat_id) if position == len(results) - 1 else None
                await _deliver(
                    chat_id,
                    lambda s=source, r=reply, m=markup: s.reply_text(
                        _plain(r), reply_markup=m, do_quote=True),
                    reply)
            # A fragment that arrived after composition passed its safe
            # point couldn't retract this reply — it starts the next round
            # now (with this reply already in context) instead of sitting
            # unprocessed until some future message wakes the flusher.
            if not _burst_buffers.get(chat_id):
                return
    finally:
        _burst_flushing.discard(chat_id)
        _burst_superseded.pop(chat_id, None)
        if typing is not None:
            typing.cancel()


_HELP_TEXT = """Hi — I'm Kyraan, your personal assistant. Talk to me like a person (voice notes work too). Things I do:

⏰ Reminders — "remind me to call mom at 7pm", "every day at 9 take medicine", "any reminders?", "cancel the call-mom one"
📅 Calendar — "what's on tomorrow?", "add lunch with Mira friday 1pm", "cancel the 3pm meeting" (I always ask before changing anything)
📬 Email — "any new emails?" (senders & subjects only — I never read bodies)
🏠 Home — "is the AC on?", "check energy", "bedroom temp", "turn off the AC"
🧠 Memory — tell me facts and I'll remember after you approve ("review memory"); "forget X" removes one
📊 Usage — "how much did we spend on AI this week?"
🌅 Briefs — mornings 7:30 and evenings 9:30, plus AC/energy heads-ups

Everything runs on your own computer; nothing is shared."""


async def _on_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """/start and /help — a bot whose very first expected message gets
    silence feels broken (completion pack)."""
    if update.effective_user is None or update.effective_chat is None:
        return
    if not _owner_private(update):
        await update.message.reply_text("I'm a private personal assistant — not open for general use.")
        return
    await update.message.reply_text(_HELP_TEXT)


async def _on_voice(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """A voice note becomes text locally (audio never leaves the Mac) and
    then flows through the exact same pipeline as a typed message."""
    if not await _media_admitted(update, "media.voice"):
        return
    from kyraan.control_plane import kernel as _idkernel
    _idkernel.set_viewer(*_viewer_for(update))  # task-scoped; PTB
    # runs each handler in its own task, so no reset is needed
    from kyraan.channels import voice

    if not await voice.wait_available():
        await update.message.reply_text(
            "I can't listen to voice notes yet on this machine — tell me in "
            "words for now.", do_quote=True)
        return
    import tempfile
    from pathlib import Path

    typing = asyncio.create_task(_typing_loop(context.bot, update.effective_chat.id))
    try:
        tg_file = await update.message.voice.get_file()
        with tempfile.NamedTemporaryFile(suffix=".oga", delete=False) as handle:
            temp_path = Path(handle.name)
        try:
            from kyraan.control_plane.logging_setup import stage as _vstage
            with _vstage("voice_transcribe"):
                await tg_file.download_to_drive(custom_path=temp_path)
                text = await voice.transcribe(temp_path)
        finally:
            temp_path.unlink(missing_ok=True)  # transient: no audio at rest
    finally:
        typing.cancel()
    if not text:
        await update.message.reply_text(
            "I couldn't make out that voice note — mind trying again, or "
            "typing it?", do_quote=True)
        return
    logger.info("Voice note transcribed (%d chars)", len(text))
    await _ingest(update, context, text)


async def _on_location(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """A shared location pin becomes text — place name via reverse
    geocoding (best-effort, coordinates as the fallback) — and flows
    through the same pipeline as a typed message, so a pin plus a caption
    like "weather here?" bursts into one thought. Live-location EDITS are
    not tracked; the initial pin is what the assistant gets. (Seen live
    2026-08-26: a shared pin matched no handler, was silently dropped,
    and the model kept asking which area the owner was in.)"""
    if not await _media_admitted(update, "media.location"):
        return
    from kyraan.control_plane import kernel as _idkernel
    _idkernel.set_viewer(*_viewer_for(update))  # task-scoped; PTB
    # runs each handler in its own task, so no reset is needed
    from kyraan.triggers import whereabouts as _where
    if update.message is None:
        # a live-location EDIT (edited_message): not a message for the
        # pipeline, but a fix for whereabouts (2026-09-04) — transitions
        # (nearly home, a known place) speak on their own
        edited = getattr(update, "edited_message", None)
        epin = getattr(edited, "location", None)
        try:
            is_owner_fix = epin is not None and _owner_private(update)
        except Exception:
            is_owner_fix = False
        if is_owner_fix:
            try:
                lines = await asyncio.to_thread(_where.observe, epin.latitude, epin.longitude)

                async def _send_where(chat_id: int, text: str) -> bool:
                    await context.bot.send_message(chat_id=chat_id, text=_plain(text))
                    orchestrator.record_proactive(chat_id, text)
                    return True
                await _where.announce(update.effective_chat.id, lines, _send_where)
            except Exception as exc:
                logger.warning("whereabouts edit failed: %s", exc)
        return
    from kyraan.channels import location as geo

    venue = getattr(update.message, "venue", None)
    pin = update.message.location or (venue.location if venue else None)
    if pin is None:
        return
    if venue and venue.title:
        # Telegram's "choose a nearby place" sends a VENUE, not a bare
        # location — it went to NO handler and vanished silently (seen
        # live 2026-08-27: a shared place produced nothing at all). The
        # venue even names itself; use that over reverse geocoding.
        described = (f"{venue.title}"
                     + (f", {venue.address}" if venue.address else "")
                     + f" ({pin.latitude:.5f}, {pin.longitude:.5f})")
    else:
        typing = asyncio.create_task(_typing_loop(context.bot, update.effective_chat.id))
        try:
            from kyraan.control_plane.logging_setup import stage as _gstage
            with _gstage("geocode_pin"):
                described = await asyncio.to_thread(geo.describe, pin.latitude, pin.longitude)
        finally:
            typing.cancel()
    logger.info("Location pin resolved: %s", described)
    if _owner_private(update):
        try:
            lines = await asyncio.to_thread(_where.observe, pin.latitude, pin.longitude)

            async def _send_where2(chat_id: int, text: str) -> bool:
                await context.bot.send_message(chat_id=chat_id, text=_plain(text))
                orchestrator.record_proactive(chat_id, text)
                return True
            await _where.announce(update.effective_chat.id, lines, _send_where2)
        except Exception as exc:
            logger.warning("whereabouts pin failed: %s", exc)
    await _ingest(update, context,
                  f"[I'm sharing my current location: {described}]")


async def _enroll_face_gated(chat_id: int, name: str, image_bytes: bytes) -> str:
    """Face enrollment through the standard confirm flow — the owner's
    "yes" (typed, or the inline button) runs the stashed enrollment with
    these exact bytes. The template is written only after that yes."""
    from kyraan.agents import faces, orchestrator
    from kyraan.control_plane.kernel import SkillCall

    if not faces.available():
        return ("Face recognition isn't set up on this machine yet — run "
                "scripts/setup_faces.py once, then send the photo again.")

    async def handler(_args):
        try:
            return await asyncio.to_thread(faces.enroll, name, image_bytes)
        except ValueError as exc:
            return f"Couldn't enroll that face: {exc}"

    return await orchestrator._gated(
        chat_id, SkillCall("faces.enroll", {"name": name}), handler,
        describe=(f'About to store a FACE TEMPLATE for "{name}" — biometric '
                  "data, kept ONLY on this machine (never sent anywhere), "
                  f'deletable anytime with "forget the face {name}"'))


async def _on_photo(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """A photo becomes one frontier vision call — analysis only, no tools
    on this path (see agents/photo.py). The reply plus a text record land
    in history so follow-up questions work."""
    if not await _media_admitted(update, "media.photo"):
        return
    from kyraan.control_plane import kernel as _idkernel
    _idkernel.set_viewer(*_viewer_for(update))  # task-scoped; PTB
    # runs each handler in its own task, so no reset is needed
    # BIOMETRICS ARE OWNER-GOVERNED (review 2026-08-28): a granted
    # media.photo means DOCUMENT/scene capture — never face operations.
    # Without this, a viewer's "its me" would have enrolled their face
    # under the OWNER'S name, and their photos would be annotated with
    # matches from the household's templates.
    is_owner_turn = _owner_private(update)
    import base64

    from kyraan.agents import orchestrator, photo
    from kyraan.control_plane.logging_setup import log_trace, new_turn

    from kyraan.agents import faces

    new_turn()
    chat_id = update.effective_chat.id
    caption = update.message.caption or ""
    log_trace("turn_start", chat_id=chat_id, user_text=f"[photo] {caption}")
    typing = asyncio.create_task(_typing_loop(context.bot, chat_id))
    try:
        # largest thumbnail Telegram offers — plenty for detail=low vision
        from kyraan.control_plane.logging_setup import stage as _stage
        with _stage("photo_download"):
            tg_file = await update.message.photo[-1].get_file()
            image_bytes = bytes(await tg_file.download_as_bytearray())
        if is_owner_turn:
            faces.stash_photo(chat_id, image_bytes)   # biometric intake is owner-only
        from kyraan.agents import secrets as _secrets
        if is_owner_turn and _secrets.active(chat_id):
            # PRIVATE MODE: photos need the cloud vision model, and the
            # promise is "everything stays on this Mac" (review 2026-09-03:
            # a photo turn went to the cloud and into history in clear).
            reply = ("🔒 Private mode is on, and reading a photo needs the "
                     "cloud vision model — so I didn't look at this one. Say "
                     "\"private mode off\" and send it again.")
            orchestrator.record_exchange(chat_id, f"[sent a photo: {caption}]", reply)
            from kyraan.control_plane.logging_setup import turn_summary
            log_trace("turn_end", chat_id=chat_id, reply=reply, **turn_summary())
            await update.message.reply_text(_plain(reply), do_quote=True)
            return

        # Either form: the strict phrase ("remember this face as X") or
        # the natural one ("remember this is Suman Ghosh") — with the
        # photo in the same message the intent is unambiguous, and the
        # confirm gate still stands (seen live 2026-08-26 23:08: the
        # natural caption described the photo instead of enrolling).
        enroll_name = faces.enroll_request(caption) if is_owner_turn else None
        if is_owner_turn and enroll_name is None:
            natural = faces.enroll_from_text(caption)
            # "remember this is Ruma's gel" on a photo with no face is a
            # naming, not an enrollment (review 2026-09-03)
            if natural and faces.face_count(image_bytes) == 1:
                enroll_name = natural
        if is_owner_turn and enroll_name is None and faces.self_claim(caption):
            # "its me": the owner asserting identity IS enrollment
            # intent for their own face — confirm gate still owns the
            # biometric write, so a "no" costs nothing.
            enroll_name = faces.owner_display_name()
        if enroll_name is None and not caption.strip():
            # The bot's own last reply may have ASKED for this photo
            # ("send a photo to save Kamal's face") — a captionless photo
            # right after is the answer, not a random picture; the
            # confirm gate below still owns the biometric write.
            recent = [t for role, t in orchestrator._history[chat_id]
                      if role == "assistant"]
            if recent and is_owner_turn:
                enroll_name = faces.invite_followup(recent[-1])
        if enroll_name is not None:
            # Biometric write → the standard confirm gate; the photo's
            # bytes stay captured in the handler for the owner's yes.
            reply = await _enroll_face_gated(chat_id, enroll_name, image_bytes)
            orchestrator.record_exchange(chat_id, f"[sent a photo: {caption}]", reply)
            from kyraan.control_plane.logging_setup import turn_summary
            log_trace("turn_end", chat_id=chat_id, reply=reply, **turn_summary())
            await update.message.reply_text(_plain(reply), do_quote=True,
                                            reply_markup=_confirm_keyboard(chat_id))
            return

        with _stage("face_recognize"):
            # Recognition consults the OWNER'S biometric templates —
            # never run for another viewer's photos.
            recognized = (await asyncio.to_thread(faces.recognize, image_bytes)
                          if faces.available() and is_owner_turn
                          else {"names": [], "maybe": [], "unknown_faces": 0})
        data_url = ("data:image/jpeg;base64,"
                    + base64.b64encode(image_bytes).decode())
        reply, vision_enroll = await photo.answer(
            chat_id, data_url, caption,
            recognized=recognized["names"],
            maybe=recognized.get("maybe") or [])
        if (vision_enroll and faces.available() and is_owner_turn
                and faces.enroll_words(caption)):
            # The vision model read enrollment intent in the caption (any
            # wording) — the regex above only catches the fixed phrases.
            # Same confirm gate; the ask replaces the descriptive reply.
            reply = await _enroll_face_gated(chat_id, vision_enroll, image_bytes)
            orchestrator.record_exchange(
                chat_id, f"[sent a photo: {caption}]", reply)
            log_trace("turn_end", chat_id=chat_id, reply=reply)
            await update.message.reply_text(
                _plain(reply), do_quote=True,
                reply_markup=_confirm_keyboard(chat_id))
            return
        if (faces.available() and is_owner_turn
                and re.search(r"who(?:'s| is)\s+(?:this|that|he|she|they|it|in\b)"
                              r"|do you (?:know|recogni[sz]e)\s+(?:him|her|them|this|who)",
                              caption, re.IGNORECASE)
                and recognized.get("unknown_faces", 0) > 0
                and not recognized["names"] and not recognized.get("maybe")):
            # A who-question with NO match answers with the truth about
            # what IS saved — live 2026-08-28: a no-match got a plain
            # visual description and the owner had to interrogate why
            # ("you said suman's face record but you can't recognise").
            enrolled = faces.enrolled_names()
            reply += ("\n\n(Checked against my saved faces — no match. "
                      + (f"I have face data for: {', '.join(enrolled)}."
                         if enrolled else "No faces are enrolled yet.")
                      + ")")
        if is_owner_turn and re.search(
                r"\b(?:similar|same|other|more|matching)\s+(?:images?|photos?|pictures?|pics?)\b",
                caption, re.IGNORECASE):
            # "similar images for kiaan? ... link it with them" (live
            # 2026-09-03 01:03): the answer is in the store, not in the
            # vision model — same-person captures with a close
            # description, and an explicit ask links them.
            from kyraan.store import documents as _docs_sim
            cap = await asyncio.to_thread(_docs_sim.latest_capture, chat_id, 1)
            sims = (await asyncio.to_thread(_docs_sim.similar_captures, cap["doc_id"])
                    if cap else [])
            if sims:
                reply += "\n\nSimilar saved photos:\n" + "\n".join(
                    f'• "{x["caption"]}" ({x["date"]})' for x in sims)
                if re.search(r"\blink", caption, re.IGNORECASE):
                    await asyncio.to_thread(_docs_sim.link_captures, cap["doc_id"],
                                            [x["doc_id"] for x in sims])
                    reply += "\n\nLinked this photo to them."
            elif cap:
                reply += "\n\nNo similar saved photos yet."
        hint_name = (faces.enroll_hint(caption)
                     if faces.available() and is_owner_turn
                     and recognized.get("unknown_faces", 0) > 0 else None)
        if hint_name:
            reply += (f'\n\n(Want me to recognize this face later? Send a solo '
                      f'photo of them captioned "remember this face as '
                      f'{hint_name}" — face data stays on this machine.)')
    except photo.VisionUnavailable:
        reply = ("I can't see photos right now (the vision model is "
                 "unavailable) — tell me in words for now.")
    except Exception as exc:
        logger.warning("Photo handling failed: %s", exc)
        reply = "Something went wrong with that photo — try sending it again."
    finally:
        typing.cancel()
    orchestrator.record_exchange(
        update.effective_chat.id,
        f"[sent a photo{': ' + caption if caption else ''}]", reply)
    from kyraan.control_plane.logging_setup import turn_summary
    log_trace("turn_end", chat_id=update.effective_chat.id, reply=reply,
              **turn_summary())
    # a photo is read by the cloud vision model — say so, every time
    for piece in _pieces(_plain(f"{reply}\n\n☁️ via cloud (vision)")):
        await update.message.reply_text(piece, do_quote=True)


_PDF_MAX_BYTES = 15 * 1024 * 1024


_TEXT_DOC_EXTS = (".txt", ".csv", ".md", ".json", ".log")


async def _on_document(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Non-PDF document intake (audit item 3, 2026-08-28): text files
    and .docx join document memory — same pipeline as PDFs (local
    extraction, original stored, hash-deduped). Anything else gets the
    honest unsupported reply."""
    if not await _media_admitted(update, "media.file"):
        return
    from kyraan.control_plane import kernel as _idkernel
    _idkernel.set_viewer(*_viewer_for(update))  # task-scoped; PTB
    # runs each handler in its own task, so no reset is needed
    document = update.message.document
    filename = (document.file_name or "").lower()
    chat_id = update.effective_chat.id
    ext = "." + filename.rsplit(".", 1)[-1] if "." in filename else ""
    if ext not in _TEXT_DOC_EXTS and ext != ".docx":
        await update.message.reply_text(
            f"I can't read {ext or 'that'} files yet — PDF, Word (.docx), "
            "or plain text/CSV work.", do_quote=True)
        return
    if (document.file_size or 0) > _PDF_MAX_BYTES:
        await update.message.reply_text(
            "That file is too large for me to keep — under "
            f"{_PDF_MAX_BYTES // (1024 * 1024)}MB works.", do_quote=True)
        return
    handle = await document.get_file()
    data = bytes(await handle.download_as_bytearray())
    if ext == ".docx":
        import io as _io

        def _read_docx() -> str:
            import docx as _docx
            parsed = _docx.Document(_io.BytesIO(data))
            parts = [p.text for p in parsed.paragraphs if p.text.strip()]
            for table in parsed.tables:
                for row in table.rows:
                    parts.append(" | ".join(c.text.strip()
                                            for c in row.cells))
            return "\n".join(parts).strip()

        try:
            text = await asyncio.to_thread(_read_docx)
        except Exception as exc:
            logger.warning("docx read failed: %s", exc)
            await update.message.reply_text(
                "I couldn't read that Word file — is it a real .docx?",
                do_quote=True)
            return
    else:
        text = data.decode("utf-8", errors="replace").strip()
    from kyraan.store import documents
    import asyncio as _aio
    doc_id = await _aio.to_thread(
        lambda: documents.ingest(
            chat_id, "file", text, (update.message.caption or "")[:120],
            document.file_name or "", original=(data, ext.lstrip(".")),
            uploaded_by=_uploader(update)))
    reply = (f'📄 Saved "{document.file_name}" to document memory '
             f"({len(text):,} chars) — ask me about it anytime." + _private_note(doc_id)
             if doc_id else
             "I couldn't distill any text worth keeping from that file.")
    await update.message.reply_text(reply, do_quote=True)
    asyncio.create_task(_enrich_and_tell(update, doc_id))
    orchestrator.record_exchange(
        chat_id, f"[sent a file: {document.file_name}]", reply)


async def _on_pdf(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Document memory (2026-08-27): a PDF from the OWNER is captured —
    text layer extracted locally (no model, no cloud), chunked, embedded
    locally, stored. Scanned PDFs without a text layer get an honest
    'can't read scans yet'. Owner-only like all media until P3.5b/c
    scope media for other stages."""
    if not await _media_admitted(update, "media.file"):
        return
    from kyraan.control_plane import kernel as _idkernel
    _idkernel.set_viewer(*_viewer_for(update))  # task-scoped; PTB
    # runs each handler in its own task, so no reset is needed
    document = update.message.document
    chat_id = update.effective_chat.id
    if (document.file_size or 0) > _PDF_MAX_BYTES:
        await update.message.reply_text(
            "That PDF is over 15MB — too big for me to take in.")
        return
    try:
        handle = await document.get_file()
        data = bytes(await handle.download_as_bytearray())
        import io

        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        text = "\n\n".join((page.extract_text() or "") for page in reader.pages).strip()
    except Exception as exc:
        logger.warning("pdf capture failed: %s", exc)
        await update.message.reply_text(
            "I couldn't read that PDF — it may be corrupted or protected.")
        return
    if len(text) < 20:
        # A scan: no text layer. OCR the embedded page images through
        # the vision tier (owner gap list 2026-08-27), first 4 pages.
        import base64

        from kyraan.agents import photo as _photo
        pieces = []
        for page in reader.pages[:4]:
            try:
                images = page.images
            except Exception:
                images = []
            for img in images[:1]:
                data_url = ("data:image/jpeg;base64,"
                            + base64.b64encode(img.data).decode())
                piece = await _photo.transcribe(data_url)
                if piece:
                    pieces.append(piece)
                break
        text = "\n\n".join(pieces).strip()
        if len(text) < 20:
            await update.message.reply_text(
                "That PDF looks scanned and I couldn't read its pages — "
                "a clear photo of the page works better.")
            return
    from kyraan.store import documents
    import asyncio as _aio
    doc_id = await _aio.to_thread(
        lambda: documents.ingest(
            chat_id, "pdf", text, (update.message.caption or "")[:120],
            document.file_name or "", original=(data, "pdf"),
            uploaded_by=_uploader(update)))
    reply = (f'📄 Saved "{document.file_name}" to document memory '
             f"({len(reader.pages)} page{'s' if len(reader.pages) != 1 else ''}, "
             f"{len(text):,} chars) — ask me about it anytime." + _private_note(doc_id)
             if doc_id else
             "I couldn't distill any text worth keeping from that PDF.")
    orchestrator.record_exchange(
        chat_id, f"[sent a PDF: {document.file_name}]", reply)
    await update.message.reply_text(reply, do_quote=True)
    asyncio.create_task(_enrich_and_tell(update, doc_id))


async def _on_unsupported(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Photos, voice notes, stickers, files — the text-only handler never
    fires for these, and the owner got SILENCE (live 2026-08-26: an image
    with 'can yiu tell me what is this' was simply ignored). Until vision
    and voice land (Phase 5), the honest reply beats a dropped message."""
    if not _owner_private(update):
        return
    message = update.message
    if message is None:
        return
    kind = ("a photo" if message.photo else
            "a voice message" if message.voice else
            "a video" if message.video else
            "an audio file" if message.audio else
            "a sticker" if message.sticker else
            "a file" if message.document else "that kind of message")
    logger.info("Unsupported media from owner: %s", kind)
    await message.reply_text(
        f"I can't open {kind} yet — I only read text for now (seeing images "
        "and hearing voice notes come in a later phase). Tell me in words "
        "and I'm all yours.",
        do_quote=True,
    )


# Misfire policy (found live 2026-08-30): APScheduler's default grace is
# ~1 second, and this Mac sleeps. A one-shot date job whose moment passed
# during sleep was DISCARDED on wake — which killed the hourly water
# series twice (its next occurrence is scheduled inside fire()), dropped
# the daily 8 PM calendar task, and swallowed Kiaan's 5 AM vaccination
# reminder outright. One-shots must fire late, never never: every fire
# path already handles lateness honestly (overdue labels, grid catch-up,
# DND holds, idempotent claims). Daily jobs get a bounded grace — a
# morning brief at 11 AM beats silence; past that it skips to tomorrow.
# Repeating pollers keep the default: skipping to the next slot IS their
# recovery.
_ALWAYS_FIRE = {"misfire_grace_time": None}
_DAILY_GRACE = {"misfire_grace_time": 4 * 3600}


async def _reminder_job(context: ContextTypes.DEFAULT_TYPE) -> None:
    data = context.job.data
    await scheduler.fire(data["reminder_id"], data["chat_id"], data["text"])


async def _agent_task_job(context: ContextTypes.DEFAULT_TYPE) -> None:
    from kyraan.triggers import agent_tasks
    data = context.job.data
    await agent_tasks.fire(data["task_id"],
                           redeliver_only=bool(data.get("redeliver_only")))


def _wire_agent_tasks(job_queue: JobQueue, bot) -> None:
    from kyraan.agents import agent_loop
    from kyraan.triggers import agent_tasks

    def schedule_fn(job_name: str, run_at, payload: dict) -> None:
        job_queue.run_once(_agent_task_job, when=run_at, data=payload,
                           name=job_name, job_kwargs=_ALWAYS_FIRE)

    async def run_fn(chat_id: int, instruction: str) -> str:
        for tier in orchestrator.tier_chain():
            try:
                return await agent_loop.run(chat_id, instruction, tier=tier, read_only=True)
            except agent_loop.AgentUnavailable:
                continue
        return ""

    async def send_fn(chat_id: int, text: str) -> None:
        if chat_id != _owner_id():
            return
        await bot.send_message(chat_id=chat_id, text=_plain(text))
        orchestrator.record_proactive(chat_id, text)

    agent_tasks.init(schedule_fn=schedule_fn, run_fn=run_fn, send_fn=send_fn,
                     only_chat=_owner_id())
    from kyraan.tools import code_agent as _code_agent
    _code_agent.init(send_fn=send_fn)   # coding-task reports come back the same way


async def _goal_cycle_job(context: ContextTypes.DEFAULT_TYPE) -> None:
    from kyraan.triggers import goals
    await goals.fire(context.job.data["goal_id"])


def _wire_goals(job_queue: JobQueue, bot) -> None:
    from kyraan.agents import agent_loop
    from kyraan.control_plane import kernel as _kernel
    from kyraan.triggers import goals

    def schedule_fn(job_name: str, run_at, payload: dict) -> None:
        for job in job_queue.get_jobs_by_name(job_name):
            job.schedule_removal()  # cadence moves: one live job per goal
        job_queue.run_once(_goal_cycle_job, when=run_at, data=payload,
                           name=job_name, job_kwargs=_ALWAYS_FIRE)

    async def run_fn(goal) -> str:
        # THE security line of goal continuity (design record 2026-08-31):
        # a cycle runs AS the goal's person — their stage bounds tool
        # reach, the §4 clause bounds fact visibility. The contextvar
        # default is owner, so an unset viewer here would hand an
        # enrolled adult's goal the owner's whole capability surface.
        token = _kernel.set_viewer(goal.person, goal.stage)
        try:
            for tier in orchestrator.tier_chain():
                try:
                    return await agent_loop.run(
                        goal.chat_id, goals.cycle_instruction(goal),
                        tier=tier, read_only=True)
                except agent_loop.AgentUnavailable:
                    continue
            return ""
        finally:
            _kernel._viewer.reset(token)

    async def send_fn(chat_id: int, text: str) -> bool:
        if chat_id != _owner_id():
            from kyraan.store import persons
            if persons.person_for_chat(chat_id, strict=True) is None:
                logger.warning("Dropping goal report for unknown chat %s",
                               chat_id)
                return False
        await bot.send_message(chat_id=chat_id, text=_plain(text))
        orchestrator.record_proactive(chat_id, text)
        return True

    goals.init(schedule_fn=schedule_fn, run_fn=run_fn, send_fn=send_fn)


def _wire_voice_echo(job_queue: JobQueue, bot) -> None:
    """Voice in the room (duty #4, 2026-09-03): poll the Echo's last
    utterance; "Alexa, Kyraan …" runs the owner's pipeline and is spoken
    back. Skips itself without Home Assistant or an allowlisted Echo."""
    import os as _os

    from kyraan.channels import voice_echo
    if not (_os.environ.get("HASS_URL") and _os.environ.get("HASS_TOKEN")):
        return
    if not voice_echo.enabled() or not voice_echo.devices():
        return

    async def _send(chat_id: int, text: str) -> bool:
        await bot.send_message(chat_id=chat_id, text=_plain(text))
        orchestrator.record_proactive(chat_id, text)
        return True

    async def _poll(context: ContextTypes.DEFAULT_TYPE) -> None:
        async with _lock_for(_owner_id()):
            await voice_echo.tick(_owner_id(), _send)

    job_queue.run_repeating(_poll, interval=voice_echo.poll_seconds(), first=5,
                            name="voice_echo", job_kwargs=_ALWAYS_FIRE)
    logger.info("Voice in the room: listening on %s every %ss",
                ", ".join(voice_echo.devices()), voice_echo.poll_seconds())


def _wire_whereabouts(job_queue: JobQueue, bot) -> None:
    """The phone's own tracker (HA person entity) every minute — silent
    until the companion app reports a position."""
    import os as _os
    if not (_os.environ.get("HASS_URL") and _os.environ.get("HASS_TOKEN")):
        return
    from kyraan.triggers import whereabouts as _where

    async def _send(chat_id: int, text: str) -> bool:
        await bot.send_message(chat_id=chat_id, text=_plain(text))
        orchestrator.record_proactive(chat_id, text)
        return True

    async def _poll(context: ContextTypes.DEFAULT_TYPE) -> None:
        async with _lock_for(_owner_id()):
            await _where.poll_person(_owner_id(), _send)

    job_queue.run_repeating(_poll, interval=60, first=20, name="whereabouts_person",
                            job_kwargs=_ALWAYS_FIRE)


def _wire_cache_warm(job_queue: JobQueue) -> None:
    """A 1-token touch of the live prompt prefix every four minutes
    while the owner is chatting, so the provider's cache stays warm."""
    from kyraan.triggers import cache_warm as _warm

    async def _tick(context: ContextTypes.DEFAULT_TYPE) -> None:
        await _warm.tick()

    job_queue.run_repeating(_tick, interval=_warm.INTERVAL_S, first=90, name="cache_warm",
                            job_kwargs=_ALWAYS_FIRE)


def _wire_slack_watch(job_queue: JobQueue, bot) -> None:
    """Mention watch (2026-09-02): draft + confirm, never post. Skips
    itself entirely when Slack isn't mounted or the token is absent."""
    import os as _os

    from kyraan.agents import agent_loop, orchestrator as _orch
    from kyraan.control_plane import kernel as _kernel
    from kyraan.control_plane.kernel import SkillCall
    from kyraan.triggers import slack_watch
    server = (_kernel.config.load().get("tool_servers") or {}).get("slack") or {}
    channels = server.get("watch_channels") or []
    token = _os.environ.get("SLACK_MCP_XOXP_TOKEN", "").strip()
    if not channels or not token:
        return

    def _whoami() -> tuple:
        # the owner's Slack identity, once, from the token itself
        import json as _json
        import urllib.request as _ur
        try:
            req = _ur.Request("https://slack.com/api/auth.test", data=b"",
                              headers={"Authorization": f"Bearer {token}"})
            body = _json.loads(_ur.urlopen(req, timeout=15).read())
            return str(body.get("user_id", "")), str(body.get("user", ""))
        except Exception as exc:
            logger.warning("slack auth.test failed: %s", exc)
            return "", ""

    _WRITER = (
        "You ghost-write ONE Slack message for Manab (the owner) to send "
        "under his own name to a family member or friend. You are not an "
        "assistant in this message — you ARE Manab typing. Use what you "
        "know about his life (below) only where it makes the reply truer. "
        "Never mention Kyraan, drafts, or assistants. Never ask the owner "
        "anything. Output the message text only.\n")

    async def draft_fn(instruction: str, question: str = "") -> str:
        # The owner-facing loop asked the OWNER for guidance and that
        # meta-talk got posted to Ruma (live 2026-09-02). Drafting is a
        # writer call: role-framed, memory-aware, no tools, no contract.
        # Memory and documents are keyed on the SENDER'S QUESTION, not
        # the whole brief — "next vaccine date" must pull the MMR fact
        # and the vaccination card, not promise to "check" (live).
        from kyraan.memory import engine as _engine
        from kyraan.model_router import router as _router
        from kyraan.store import documents as _docs
        key = question or instruction
        memory, docs = "", ""
        informational = slack_watch.is_informational(question) if question else True
        try:
            if informational:
                memory = _engine.build_context(key, budget_chars=1800)
        except Exception:
            memory = ""
        try:
            hits = (await asyncio.to_thread(_docs.search, _owner_id(), key)
                    if informational else [])
            docs = "\n".join(f'- [{h["caption"]}, {h["date"]}] {h["text"][:300]}'
                              for h in (hits or [])[:3])
        except Exception:
            docs = ""
        system = (_WRITER
                  + (f"\nWHAT MANAB KNOWS (facts):\n{memory}" if memory else "")
                  + (f"\nMANAB'S DOCUMENTS:\n{docs}" if docs else ""))
        for tier in orchestrator.tier_chain():
            try:
                resp = await _router.acall(prompt=instruction, system=system,
                                           tier=tier, max_tokens=300)
                return (resp.text or "").strip().strip('"')
            except Exception:
                continue
        return ""

    async def ask_fn(chat_id: int, channel: str, draft: str, context: str) -> None:
        if not draft:
            # two rejected drafts: the mention still reaches the owner,
            # honestly without a proposal
            text = (f"📣 Slack {channel} — {context}\n\n(I couldn't write a "
                    "natural reply for this one — answer in Slack yourself.)")
            await bot.send_message(chat_id=chat_id, text=_plain(text))
            _orch.record_proactive(chat_id, text)
            return

        async def _post(_a: dict) -> str:
            result = await _kernel.run_tool(_kernel.ToolCall(
                "slack.post", {"channel_id": channel, "payload": draft,
                               "content_type": "text/plain"}))
            slack_watch.note_posted(draft)   # never a voice sample
            return f"Posted to {channel}: \"{draft}\""
        ask = await _orch._gated(
            chat_id, SkillCall("agent.action", {"tool": "slack.post"}), _post,
            describe=(f"📣 Slack {channel} — {context}\n\n"
                      f"Proposed reply (posted as you): \"{draft}\""))
        await bot.send_message(chat_id=chat_id, text=_plain(ask),
                               reply_markup=_confirm_keyboard(chat_id))
        _orch.record_proactive(chat_id, ask)

    async def _job(context: ContextTypes.DEFAULT_TYPE) -> None:
        try:
            if not slack_watch._owner_user_id:
                uid, handle = await asyncio.to_thread(_whoami)
                if not uid:
                    return
                slack_watch.init(uid, draft_fn, ask_fn, owner_handle=handle)
            await slack_watch.tick(channels, _owner_id())
        except Exception as exc:
            logger.warning("slack watch tick failed: %s", exc)

    job_queue.run_repeating(_job, interval=120, first=90, name="slack_watch")
    logger.info("Slack mention watch armed (%d channels, every 2 min)", len(channels))


def _wire_scheduler(job_queue: JobQueue, bot) -> None:
    def schedule_fn(job_name: str, run_at, payload: dict) -> None:
        job_queue.run_once(_reminder_job, when=run_at, data=payload,
                           name=job_name, job_kwargs=_ALWAYS_FIRE)

    def cancel_fn(job_name: str) -> None:
        for job in job_queue.get_jobs_by_name(job_name):
            job.schedule_removal()

    async def send_fn(chat_id: int, text: str) -> bool:
        # The store is shared with the dev harnesses (chat.py uses chat 0,
        # walkthrough scripts use their own ids). A record like that would
        # make send_message error on a nonexistent chat and leave the
        # reminder pending forever, retried on every restart. Deliverable
        # chats are the owner AND admitted enrolled persons (Bugbot P1,
        # 2026-08-28: reminders.create is in every viewer stage's toolset,
        # but this owner-era gate silently retired a viewer's reminders —
        # Ruma could set one and never receive it); anything else is
        # retired, and the False return makes fire() log it truthfully as
        # reminder_retired_undelivered, not reminder_sent.
        if chat_id != _owner_id():
            from kyraan.store import persons
            # strict: a PG outage must RAISE (fire() retries) — returning
            # False here permanently retires an enrolled person's
            # reminder over a transient outage (Bugbot round-2 P1).
            if persons.person_for_chat(chat_id, strict=True) is None:
                logger.warning("Retiring reminder for unknown chat %s "
                               "(dev-harness record)", chat_id)
                return False
        await bot.send_message(chat_id=chat_id, text=text)
        orchestrator.record_proactive(chat_id, text)
        if chat_id == _owner_id():
            # Spoken reminders (owner decision 2026-09-02): the OWNER's
            # reminders also arrive as voice on the Echo. Best-effort and
            # contained — a TTS failure never fails the reminder; fire()
            # already holds DND, so a spoken reminder can't break quiet
            # hours. Enrolled persons stay Telegram-only.
            try:
                from kyraan.tools import home_assistant as _ha
                if _ha._announce_targets():
                    import asyncio as _aio5
                    await _aio5.to_thread(_ha._announce, text[:240])
            except Exception as exc:
                logger.debug("spoken reminder skipped: %s", exc)
        return True

    scheduler.init(schedule_fn=schedule_fn, cancel_fn=cancel_fn, send_fn=send_fn)


def _wire_brief(job_queue: JobQueue, bot) -> None:
    async def _send(context, chat_id: int, text: str) -> bool:
        # Proactive sends get the same delivery truth as replies
        # (audit nit, 2026-08-28): one retry, then a recorded
        # reply_delivery_failed carrying the undelivered text. The
        # bool PROPAGATES (Bugbot P1: discarding it let a failed brief/
        # alert be marked sent and permanently suppressed).
        async def _once():
            await context.bot.send_message(chat_id=chat_id, text=text)

        ok = await _deliver(chat_id, _once, text)
        if ok:
            orchestrator.record_proactive(chat_id, text)   # never inside the retried send
        return ok

    at = briefs.brief_time("morning")
    if at is not None:
        async def _morning_job(context: ContextTypes.DEFAULT_TYPE) -> None:
            await briefs.fire(_owner_id(), lambda c, t: _send(context, c, t))

        # run_daily needs a tz-aware time or it fires in UTC.
        job_queue.run_daily(_morning_job, time=at.replace(tzinfo=local_now().tzinfo),
                            name="morning_brief", job_kwargs=_DAILY_GRACE)
        logger.info("Morning brief scheduled daily at %s %s", at, local_now().tzinfo)

    evening_at = briefs.brief_time("evening")
    if evening_at is not None:
        async def _evening_job(context: ContextTypes.DEFAULT_TYPE) -> None:
            await briefs.fire_evening(_owner_id(), lambda c, t: _send(context, c, t))

        job_queue.run_daily(_evening_job,
                            time=evening_at.replace(tzinfo=local_now().tzinfo),
                            name="evening_brief", job_kwargs=_DAILY_GRACE)
        logger.info("Evening brief scheduled daily at %s", evening_at)

    # Duty #1 — Kiaan's keeper (2026-09-03): the morning check, after the
    # brief, same proactive gate and delivery truth.
    from kyraan.triggers import kiaan_keeper as _keeper
    keeper_at = _keeper.check_time()
    if keeper_at is not None:
        async def _keeper_job(context: ContextTypes.DEFAULT_TYPE) -> None:
            await _keeper.fire(_owner_id(), lambda c, t: _send(context, c, t))

        job_queue.run_daily(_keeper_job, time=keeper_at.replace(tzinfo=local_now().tzinfo),
                            name="kiaan_keeper", job_kwargs=_DAILY_GRACE)
        logger.info("Kiaan's keeper scheduled daily at %s", keeper_at)

    # Duty #3 — chief of staff (2026-09-03): the 18:00 "still open".
    from kyraan.triggers import chief_of_staff as _cos
    cos_at = _cos.still_open_time()
    if cos_at is not None:
        async def _still_open_job(context: ContextTypes.DEFAULT_TYPE) -> None:
            await _cos.fire_still_open(_owner_id(), lambda c, t: _send(context, c, t))

        job_queue.run_daily(_still_open_job, time=cos_at.replace(tzinfo=local_now().tzinfo),
                            name="still_open", job_kwargs=_DAILY_GRACE)
        logger.info("Chief of staff 'still open' scheduled daily at %s", cos_at)

    # Duty #2 — house steward (2026-09-03): the 21:45 settle check.
    from kyraan.triggers import house_steward as _steward
    settle_at = _steward.settle_time()
    if settle_at is not None:
        async def _settle_job(context: ContextTypes.DEFAULT_TYPE) -> None:
            await _steward.fire_settle(_owner_id(), lambda c, t: _send(context, c, t))

        job_queue.run_daily(_settle_job, time=settle_at.replace(tzinfo=local_now().tzinfo),
                            name="house_settle", job_kwargs=_DAILY_GRACE)
        logger.info("House steward settle check scheduled daily at %s", settle_at)

    review_at = __import__("kyraan.triggers.self_review", fromlist=["x"]).review_time()
    if review_at is not None:
        from kyraan.triggers import self_review

        async def _review_job(context: ContextTypes.DEFAULT_TYPE) -> None:
            # Every nightly stage runs INDEPENDENTLY. They were chained,
            # so a failure in the critique or the episode ingest skipped
            # everything after it — including the forget re-sweep, which
            # is what keeps forgotten topics from lingering findable in
            # new episodes. A privacy repair must not depend on an
            # unrelated step succeeding (Bugbot P1).
            import asyncio as _aio
            from datetime import timedelta as _td

            from kyraan.control_plane.dnd import local_now as _now

            async def _stage(name: str, run):
                try:
                    await run()
                except Exception as exc:
                    logger.exception("nightly stage %s failed", name)
                    log_event("nightly_stage_failed", stage=name,
                              error=str(exc)[:200],
                              error_type=type(exc).__name__)

            await _stage("self_review", lambda: self_review.fire(
                _owner_id(), lambda c, t: _send(context, c, t)))

            # P3.3b: the nightly episode write rides the same job —
            # yesterday+today because ingest is idempotent and last
            # night's run stopped at this hour. Blocking work (local
            # embed + tagging + pg) stays off the event loop.
            async def _ingest():
                from kyraan.store import episodes as _episodes
                today = _now().date()
                days = [(today - _td(days=1)).isoformat(), today.isoformat()]
                await _aio.to_thread(_episodes.ingest_recent, days)

            await _stage("episode_ingest", _ingest)

            # The self-heals ride along nightly (gap audit 2026-08-27:
            # they only ran on a manual resync). Separate stages: the
            # forget re-sweep is a privacy repair and the graph catch-up
            # is bookkeeping — neither may block the other.
            async def _resweep():
                from kyraan.memory import engine as _engine
                await _aio.to_thread(_engine.resweep_forgotten)

            async def _graph_catch_up():
                from kyraan.store import triples as _triples
                await _aio.to_thread(_triples.catch_up)

            await _stage("forget_resweep", _resweep)

            # Orphaned originals (Bugbot round-3 P2): files whose doc
            # row is gone — failed unlinks, crash-in-the-gap leftovers.
            async def _orphan_sweep():
                from kyraan.store import documents as _documents
                await _aio.to_thread(_documents.sweep_orphaned_files)

            await _stage("document_orphan_sweep", _orphan_sweep)
            await _stage("graph_catch_up", _graph_catch_up)

            # P3.5e: expired 24h objection windows promote nightly.
            async def _auto_approvals():
                from kyraan.memory import review_scaling as _scaling
                await _aio.to_thread(_scaling.sweep_auto_approvals)

            await _stage("auto_approvals", _auto_approvals)

            # Cross-person contradiction scan (multi-user audit fix):
            # undeclared conflicts become the standard dispute state —
            # both facts flagged, a resolvable notice in the right queue.
            async def _conflict_scan():
                from kyraan.control_plane import kernel as _kernel
                from kyraan.memory import conflicts as _conflicts
                filed = await _aio.to_thread(_conflicts.nightly_scan)
                if filed and _kernel.can_send_proactively():
                    await _send(context, _owner_id(),
                                f"⚖️ Memory: {filed} cross-person "
                                "contradiction(s) found — both versions "
                                'stand as disputed; say "review memory" '
                                "to resolve.")

            await _stage("conflict_scan", _conflict_scan)

            # The nightly doctor: silent when healthy; WARN/FAIL sends
            # the owner the needs-work list (DND-gated like every
            # proactive).
            async def _health_report():
                from kyraan.control_plane import health as _health
                from kyraan.control_plane import kernel as _kernel
                verdict, text = await _aio.to_thread(_health.report)
                if verdict != "OK" and _kernel.can_send_proactively():
                    await _send(context, _owner_id(),
                                f"🩺 Nightly health: {verdict}\n{text}")

            await _stage("health_report", _health_report)

            # Semantic dedup scan (owner: "make it automate",
            # 2026-08-27): the model PROPOSES nightly; applying stays
            # behind the owner's yes via the "consolidate memory" chat
            # phrase. DND-gated like every proactive.
            async def _dedup_scan():
                from kyraan.control_plane import kernel as _kernel
                from kyraan.memory import consolidate as _consolidate
                proposals = await _aio.to_thread(_consolidate.scan)
                if proposals and _kernel.can_send_proactively():
                    lines = [f'• keep "{p["keep_content"][:70]}" over '
                             + "; ".join(f'"{c[:60]}"' for _, c in p["duplicates"])
                             for p in proposals]
                    await _send(context, _owner_id(),
                                "🧹 Memory dedup: "
                                f"{len(proposals)} duplicate group(s) found:\n"
                                + "\n".join(lines)
                                + '\n\nSay "consolidate memory" to review and apply.')

            await _stage("memory_dedup_scan", _dedup_scan)

            async def _lesson_scan():
                # Correction→behavior loop (2026-09-01): cluster the
                # owner's repeated corrections, draft ONE rule on the
                # LOCAL tier, queue it for review — nothing changes
                # without the owner's yes.
                from kyraan.memory import lessons
                proposed = await lessons.scan_and_propose()
                if proposed:
                    await _send(context, _owner_id(),
                                f"📐 I keep getting corrected the same way — "
                                f"{proposed} proposed behavior rule(s) queued. "
                                'Say "review memory" to see and decide.')

            await _stage("lesson_scan", _lesson_scan)

            async def _contacts_sync():
                # Nightly-job-only by governance precondition (plan §3c):
                # sync is never an agent-callable tool.
                from kyraan.store import contacts as _cstore
                from kyraan.tools import google_contacts as _gc
                if not _gc.enabled():
                    return
                import asyncio as _aio4
                fetched = await _aio4.to_thread(_gc.fetch_all)
                await _aio4.to_thread(_cstore.upsert_all, fetched)

            await _stage("contacts_sync", _contacts_sync)

            async def _vault_sync():
                from kyraan.store import notes as _notes
                if _notes.vault_root() is None:
                    return
                import asyncio as _aio6
                await _aio6.to_thread(_notes.sync, _owner_id())

            await _stage("vault_sync", _vault_sync)

        job_queue.run_daily(_review_job,
                            time=review_at.replace(tzinfo=local_now().tzinfo),
                            name="self_review", job_kwargs=_DAILY_GRACE)
        logger.info("Nightly self-review scheduled at %s", review_at)

    from kyraan.triggers import home_alerts
    if home_alerts.enabled():
        async def _alerts_job(context: ContextTypes.DEFAULT_TYPE) -> None:
            await home_alerts.check(_owner_id(), lambda c, t: _send(context, c, t))

        job_queue.run_repeating(_alerts_job, interval=1800, first=120,
                                name="home_alerts")
        logger.info("Home alerts armed (every 30 min, DND-gated)")

    # Event-triggered watch rules (owner, 2026-08-27): user-defined
    # conditions evaluated every 15 min — notify-only by doctrine.
    from kyraan.triggers import event_rules

    # Same-day episode catch-up (2026-08-28): recall stays within ~30
    # min of live instead of waiting for the 21:45 ingest.
    async def _episode_catchup(context: ContextTypes.DEFAULT_TYPE) -> None:
        import asyncio as _aio2
        from kyraan.store import episodes as _episodes
        try:
            await _aio2.to_thread(_episodes.catch_up_today)
        except Exception as exc:
            logger.warning("episode catch-up failed: %s", exc)

    job_queue.run_repeating(_episode_catchup, interval=1800, first=600,
                            name="episode_catchup")

    async def _rules_job(context: ContextTypes.DEFAULT_TYPE) -> None:
        # Send via the job's context like every proactive job here — a
        # module-level send_fn captured at wiring time referenced a name
        # not in this scope and every tick died with NameError (found
        # live: the owner's first rule never fired at 27.4°C).
        await event_rules.tick(lambda c, t: _send(context, c, t))

    job_queue.run_repeating(_rules_job, interval=900, first=180,
                            name="event_rules")

    # Wake planner (§3d #4, 2026-08-31): keep ONE pmset wake armed for
    # the next due moment so sleep can't delay what the misfire fix
    # already keeps from being lost. Thread offload — pmset is a
    # subprocess and the tick must never block the loop.
    async def _wake_job(context: ContextTypes.DEFAULT_TYPE) -> None:
        import asyncio as _aio3

        from kyraan.control_plane import wake
        try:
            await _aio3.to_thread(wake.plan)
        except Exception as exc:
            logger.warning("wake planner failed: %s", exc)

    job_queue.run_repeating(_wake_job, interval=900, first=60,
                            name="wake_planner")
    logger.info("Wake planner armed (every 15 min)")
    logger.info("Watch rules armed (every 15 min, DND-gated)")


def _harden_data_permissions() -> None:
    """Personal data is owner-only on disk (security round P2: 0644/0755
    let any local account read the chat log and memory tree)."""
    import stat
    from pathlib import Path

    repo = Path(__file__).resolve().parents[3]
    for name in ("logs", "memory", "data"):
        root = repo / name
        if not root.exists():
            continue
        os.chmod(root, 0o700)
        for path in root.rglob("*"):
            try:
                os.chmod(path, 0o700 if path.is_dir() else 0o600)
            except OSError:
                pass


def _validate_startup() -> None:
    """Fail LOUDLY at boot instead of latently at first use (review P2):
    the tool registry's load-time validation (unknown servers, auto
    writes, bad fallbacks) and the model-tier wiring both run now."""
    from kyraan.control_plane import config
    from kyraan.tools import registry

    int(os.environ["TELEGRAM_OWNER_ID"])  # KeyError/ValueError = boot failure

    tools = registry.load()
    cfg = config.load()
    tiers = cfg["model_tiers"]
    providers = cfg["providers"]
    for required in ("cheap", "frontier"):
        if required not in tiers:
            raise ValueError(f"model_tiers must define {required!r} — the fallback chain depends on it")
    _KNOWN_KINDS = {"anthropic", "gemini", "openai_compatible", "ollama_native"}
    for pname, provider in providers.items():
        kind = provider.get("kind")
        if kind not in _KNOWN_KINDS:
            raise ValueError(f"provider {pname!r} has unknown kind {kind!r}")
        if kind in ("anthropic", "gemini") and not provider.get("api_key_env"):
            raise ValueError(f"provider {pname!r} ({kind}) declares no api_key_env")
        if (kind in ("openai_compatible", "ollama_native")
                and not provider.get("api_key_env")
                and provider.get("allow_unauthenticated") is not True):
            # Keyless must be DECLARED, not inferred from the hostname
            # (round-6 P2: a local authenticated proxy broke the implicit
            # localhost-is-keyless assumption silently).
            # `is True` on purpose: "false" (string) and 1 are truthy —
            # a mistyped security flag must not grant the bypass (round-7).
            raise ValueError(
                f"provider {pname!r} has no api_key_env — if that is intentional "
                "(local unauthenticated server), set allow_unauthenticated: true "
                "(Boolean true exactly)")
    for name, tier in tiers.items():
        provider = providers.get(tier.get("provider"))
        if provider is None:
            raise ValueError(f"model tier {name!r} names unknown provider {tier.get('provider')!r}")
        if not tier.get("model"):
            raise ValueError(f"model tier {name!r} has no model")
        key_env = provider.get("api_key_env")
        if key_env and not os.environ.get(key_env, "").strip():
            raise ValueError(
                f"model tier {name!r} needs {key_env} in .env (provider {tier['provider']!r})")

    import importlib
    for server_name, server in (cfg.get("tool_servers") or {}).items():
        if server.get("transport") == "builtin" and server.get("module"):
            importlib.import_module(server["module"])  # ImportError = boot failure

    logger.info("Startup validation: %d tools, %d model tiers, owner id OK",
                len(tools), len(tiers))


def run() -> None:
    _validate_startup()
    _harden_data_permissions()
    from kyraan.channels import voice as _voice_probe
    _voice_probe.start_probe()  # verdict ready before the first voice note
    token = os.environ["TELEGRAM_BOT_TOKEN"]
    from telegram import LinkPreviewOptions
    from telegram.ext import Defaults
    app = (Application.builder().token(token).concurrent_updates(True)
           # No link-preview cards (owner: "this is not convenient",
           # 2026-08-31 — one headline reply ballooned into a giant
           # site banner). URLs stay clickable; the chat stays a chat.
           .defaults(Defaults(
               link_preview_options=LinkPreviewOptions(is_disabled=True)))
           .build())
    app.add_handler(CommandHandler(["start", "help"], _on_command))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, _on_message))
    app.add_handler(MessageHandler(filters.VOICE, _on_voice))
    app.add_handler(MessageHandler(filters.LOCATION | filters.VENUE, _on_location))
    app.add_handler(MessageHandler(filters.PHOTO, _on_photo))
    app.add_handler(MessageHandler(filters.Document.PDF, _on_pdf))
    app.add_handler(MessageHandler(filters.Document.ALL & ~filters.Document.PDF,
                                   _on_document))
    app.add_handler(MessageHandler(
        filters.VIDEO | filters.AUDIO
        | filters.Sticker.ALL | filters.Document.ALL,
        _on_unsupported,
    ))
    app.add_handler(CallbackQueryHandler(_on_callback, pattern="^kyraan_(yes|no)"))

    _wire_scheduler(app.job_queue, app.bot)
    _wire_agent_tasks(app.job_queue, app.bot)
    _wire_goals(app.job_queue, app.bot)
    _wire_slack_watch(app.job_queue, app.bot)
    _wire_brief(app.job_queue, app.bot)
    _wire_voice_echo(app.job_queue, app.bot)
    _wire_whereabouts(app.job_queue, app.bot)
    _wire_cache_warm(app.job_queue)

    # Files OUT (2026-08-28): the loop's files.send delivers through here.
    from kyraan.channels import file_send as _file_send

    async def _send_document(chat_id: int, filename: str, data: bytes,
                             caption: str) -> None:
        import io
        if filename.lower().endswith((".jpg", ".jpeg", ".png")):
            # An image DISPLAYS inline (owner verify 2026-08-28 01:54:
            # the supplement photo came back as an attachment card, not
            # a visible picture). Telegram re-compresses photos; the
            # pristine original stays in data/documents either way.
            await app.bot.send_photo(chat_id=chat_id,
                                     photo=io.BytesIO(data),
                                     caption=caption or filename)
            return
        await app.bot.send_document(
            chat_id=chat_id, document=io.BytesIO(data), filename=filename,
            caption=caption or None)

    _file_send.init(_send_document)
    # A restart must be invisible to the owner: reload the conversation
    # from chat.jsonl so follow-ups ("are those the latest emails?") still
    # have their context.
    orchestrator.seed_history_from_log()
    from kyraan.memory import engine
    engine.migrate_from_tree()  # one-time index backfill; no-op after

    logger.info("Kyraan Telegram bot starting (owner-only, Phase 1)")
    app.run_polling()
